import os
import json
import yaml
import argparse
from pathlib import Path
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import openai
import groq
from tenacity import retry, stop_after_attempt, wait_exponential
import tiktoken

# Load configuration
with open("config.yaml") as f:
    config = yaml.safe_load(f)

load_dotenv()


class DocumentProcessor:
    """Handles document input/output operations with format preservation"""

    SUPPORTED_FORMATS = [".pdf", ".docx", ".txt"]

    def __init__(self):
        self.encoder = tiktoken.get_encoding("cl100k_base")

    def read_file(self, file_path):
        self._validate_file(file_path)
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._read_pdf(file_path)
        elif ext == ".docx":
            return self._read_docx(file_path)
        elif ext == ".txt":
            return self._read_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def write_file(self, content, output_path):
        ext = Path(output_path).suffix.lower()

        if ext == ".docx":
            self._write_docx(content, output_path)
        else:
            with open(output_path, "w") as f:
                f.write(content)

    def _validate_file(self, file_path):
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if Path(file_path).suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {Path(file_path).name}")

    def _read_pdf(self, file_path):
        text = ""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        return text

    def _read_docx(self, file_path):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def _read_txt(self, file_path):
        with open(file_path, "r") as f:
            return f.read()

    def _write_docx(self, content, output_path):
        doc = Document()
        for paragraph in content.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(output_path)

    def count_tokens(self, text):
        return len(self.encoder.encode(text))


class LLMClient:
    """Unified client for OpenAI and Groq"""

    def __init__(self, provider="openai"):
        self.provider = provider
        if provider == "openai":
            self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        elif provider == "groq":
            self.client = groq.Groq(api_key=os.getenv("GROQ_API_KEY"))
        else:
            raise ValueError("Unsupported LLM provider")

    @retry(
        stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    def chat_completion(self, model, messages, **kwargs):
        if self.provider == "openai":
            return self.client.chat.completions.create(
                model=model, messages=messages, **kwargs
            )
        elif self.provider == "groq":
            return self.client.chat.completions.create(
                model=model,
                messages=messages,
                **{
                    k: v
                    for k, v in kwargs.items()
                    if k in ["temperature", "max_tokens", "response_format"]
                },
            )


class JobAnalyzer:
    """Performs deep analysis of job descriptions"""

    def __init__(self, llm_client):
        self.llm = llm_client

    def analyze(self, jd_text):
        prompt = f"""Analyze this job description with extreme attention to detail:
        {jd_text[:6000]}

        Return JSON with:
        - "hard_requirements": List of 5-7 absolute must-have qualifications
        - "preferred_skills": List of 5 nice-to-have skills
        - "success_metrics": List of measurable outcomes expected in first year
        - "cultural_fit": List of 3 inferred company values
        - "role_keywords": List of 10 critical keywords/phrases
        - "reporting_level": Inferred seniority (entry, mid, senior, exec)
        - "tech_stack": List of mentioned or implied technologies"""

        response = self.llm.chat_completion(
            model=config["models"][self.llm.provider]["job_analyzer"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        return json.loads(response.choices[0].message.content)


class CVInsightExtractor:
    """Extracts key insights from existing CV"""

    def __init__(self, llm_client):
        self.llm = llm_client

    def extract(self, cv_text):
        prompt = f"""Analyze this CV and extract:
        - "core_competencies": Top 3 technical skills
        - "key_achievements": Most impressive accomplishments
        - "career_trajectory": Progression pattern
        - "unique_value": Unique selling proposition
        - "skill_gaps": Missing requirements for target roles

        {cv_text[:4000]}"""

        response = self.llm.chat_completion(
            model=config["models"][self.llm.provider]["cv_insights"],
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(response.choices[0].message.content)


class CVCustomizer:
    """Optimizes CV for specific job applications"""

    def __init__(self, llm_client):
        self.llm = llm_client

    def customize(self, cv_text, analysis):
        prompt = f"""Transform this CV into a job-winning document:
        
        # Job Requirements
        {json.dumps(analysis, indent=2)}
        
        # Original CV
        {cv_text[:4000]}
        
        # Customization Rules
        1. Professional Summary: Mirror top 3 hard_requirements
        2. Experience: For each role:
           - Keep 2 original bullets
           - Add 1 new bullet using {analysis["success_metrics"][0]}
           - Start with {analysis["role_keywords"][0]} verbs
        3. Skills: Order by job's hard_requirements
        4. Projects: Add 1 project showing 2 preferred_skills
        5. Use {analysis["tech_stack"]} keywords naturally
        6. ATS Optimization: Simple headers, no tables/graphics"""

        response = self.llm.chat_completion(
            model=config["models"][self.llm.provider]["cv_customizer"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        return response.choices[0].message.content


class CoverLetterCustomizer:
    """Creates targeted cover letters"""

    def __init__(self, llm_client):
        self.llm = llm_client

    def customize(self, cl_text, analysis, cv_insights):
        prompt = f"""Write a compelling cover letter using this strategy:
        
        # Job Analysis
        {json.dumps(analysis, indent=2)}
        
        # Applicant Profile
        {json.dumps(cv_insights, indent=2)}
        
        # Writing Framework
        1. Opening: "As a [role] with [X] years experience in [Y], I..."
        2. Body 1: Address {analysis["hard_requirements"][0]} with CAR example
        3. Body 2: Connect {analysis["cultural_fit"][0]} to work philosophy
        4. Body 3: Align {cv_insights["unique_value"]} with {analysis["success_metrics"][0]}
        5. Closing: Clear call-to-action
        6. Style: {analysis["cultural_fit"][1]} tone
        7. Length: 3 paragraphs (250-350 words)"""

        response = self.llm.chat_completion(
            model=config["models"][self.llm.provider]["cl_customizer"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return response.choices[0].message.content


class ApplicationBuilder:
    """Main application orchestration"""

    def __init__(self, provider="openai"):
        self.processor = DocumentProcessor()
        self.llm = LLMClient(provider)
        self.analyzer = JobAnalyzer(self.llm)
        self.cv_insights = CVInsightExtractor(self.llm)
        self.cv_customizer = CVCustomizer(self.llm)
        self.cl_customizer = CoverLetterCustomizer(self.llm)

    def build(self, cv_path, cl_path, jd_path, output_dir):
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)

        # Read documents
        cv_text = self.processor.read_file(cv_path)
        cl_text = self.processor.read_file(cl_path)
        jd_text = self.processor.read_file(jd_path)

        # Analyze inputs
        job_analysis = self.analyzer.analyze(jd_text)
        cv_analysis = self.cv_insights.extract(cv_text)

        # Customize documents
        customized_cv = self.cv_customizer.customize(cv_text, job_analysis)
        customized_cl = self.cl_customizer.customize(cl_text, job_analysis, cv_analysis)

        # Save outputs
        cv_output = output_dir / f"Optimized_CV_{Path(cv_path).stem}.docx"
        cl_output = output_dir / f"Custom_CL_{Path(cl_path).stem}.docx"
        self.processor.write_file(customized_cv, cv_output)
        self.processor.write_file(customized_cl, cl_output)

        return {
            "cv_path": str(cv_output),
            "cl_path": str(cl_output),
            "analysis": job_analysis,
        }


class ApplicationCLI:
    """Command line interface handler"""

    def __init__(self):
        self.provider = config.get("provider", "openai")

    def run(self):
        parser = argparse.ArgumentParser(
            description="AI-Powered Job Application Builder"
        )
        parser.add_argument("--cv", required=True, help="Path to CV file")
        parser.add_argument("--cl", required=True, help="Path to Cover Letter file")
        parser.add_argument("--jd", required=True, help="Path to Job Description file")
        parser.add_argument("--output", default="output", help="Output directory")
        parser.add_argument(
            "--provider",
            choices=["openai", "groq"],
            default=self.provider,
            help="LLM provider",
        )
        args = parser.parse_args()

        try:
            app = ApplicationBuilder(args.provider)
            result = app.build(args.cv, args.cl, args.jd, args.output)

            print("\n✅ Application Package Created:")
            print(f"📄 CV: {result['cv_path']}")
            print(f"📝 Cover Letter: {result['cl_path']}")
            print("\n🔍 Job Analysis Insights:")
            print(
                f"- Must-Have: {', '.join(result['analysis']['hard_requirements'][:3])}..."
            )
            print(f"- Success Metrics: {result['analysis']['success_metrics'][0]}")
            print(f"- Tech Stack: {', '.join(result['analysis']['tech_stack'][:3])}")

        except Exception as e:
            print(f"❌ Error: {str(e)}")
            exit(1)


if __name__ == "__main__":
    ApplicationCLI().run()
